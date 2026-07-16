from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("output/docs/AIFX_Phase2_Face_Enhancement_API_4_Day_Plan_EN.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 32, 46)
MUTED = RGBColor(90, 96, 106)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.08


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_run_font(run, size=None, color=None, bold=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def add_para(doc, text="", style=None, size=11, color=INK, bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)
        r = p.add_run(line)
        set_run_font(r, size=9.3, color=INK, name="Courier New")


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_margins(table, 120, 160, 120, 160)
    set_table_borders(table, "C8D5E2")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    add_para(doc, "", after=6)


def add_simple_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_cell_margins(table)
    set_table_borders(table)
    set_repeat_header(table.rows[0])
    hdr = table.rows[0].cells
    for idx, label in enumerate(headers):
        set_cell_shading(hdr[idx], header_fill)
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=font_size, color=INK, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(value)
            set_run_font(r, size=font_size, color=INK)
    return table


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = "AIFX Studio Phase 2 | ComfyUI API Integration Plan"
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Confidential working plan"
    set_run_font(footer.runs[0], size=9, color=MUTED)

    add_para(doc, "AIFX Studio Phase 2", size=13, color=MUTED, bold=True, after=4)
    add_para(doc, "Face Enhancement API Four-Day Delivery Plan", size=24, color=INK, bold=True, after=4)
    add_para(
        doc,
        "PDF-aligned implementation plan for ComfyUI API integration, LoRA mapping, spatial merging, and validation.",
        size=12,
        color=MUTED,
        after=14,
    )

    meta_rows = [
        ("Target application", "AIFX Studio Integration"),
        ("Phase 2 deadline", "July 13, 2026"),
        ("Prepared by", "Zooey Chen"),
        ("Date", "2026-07-07"),
        ("Workflow template", "ComfyUI API prompt template: zooey.json"),
        ("Immediate test case", "Submit one cropped face image to ComfyUI and return one enhanced cropped face image"),
    ]
    for label, value in meta_rows:
        p = add_para(doc, "", after=2)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=10.5, color=INK, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=10.5, color=INK)

    add_callout(
        doc,
        "Executive Summary",
        "This plan matches the Phase 2 assignment. The implementation will extend the Phase 1 face detection output, route cropped face regions through a ComfyUI API workflow with character-specific LoRAs, and merge enhanced faces back into the original high-resolution image. The first practical test will use one crop image to validate the ComfyUI API loop before expanding to multi-face processing and final spatial blending.",
    )

    add_heading(doc, "1. Assignment Alignment", 1)
    alignment_rows = [
        ("ComfyUI API wrapper", "Covered", "Use HTTP/WebSocket API, /prompt submission, history polling, and output retrieval."),
        ("Dynamic JSON construction", "Covered", "Use zooey.json as a fixed template and inject image, LoRA, prompt, and output prefix at runtime."),
        ("Character-to-LoRA mapping", "Covered", "Create lora_config.json or a table mapping character_id/name to exact LoRA filenames."),
        ("Multi-character handling", "Covered", "Loop through each detected face and call ComfyUI sequentially first, with parallel execution as an optimization."),
        ("Spatial merging and blending", "Covered", "Use Phase 1 bbox metadata and blend the enhanced crop using alpha feathering, mask blur, or Poisson blending."),
        ("Validation UI", "Covered", "Extend the Phase 1 UI with input, character assignment, logs, and visual comparison."),
        ("Standalone API endpoint", "Covered", "Provide one protected POST endpoint for the automated Phase 1 + Phase 2 chain."),
    ]
    add_simple_table(doc, ["PDF Requirement", "Status", "Plan Coverage"], alignment_rows, [2450, 1200, 5710], LIGHT_BLUE, 8.9)

    add_heading(doc, "2. Project Objective", 1)
    add_para(
        doc,
        "The Phase 2 objective is to turn the Phase 1 face detection prototype into an API-driven face enhancement pipeline. The service will receive an original image or detected crop data, identify the correct character LoRA, submit each cropped face region to ComfyUI, retrieve the enhanced face output, and merge the result back into the original image without visible bounding-box edges.",
    )
    add_bullet(doc, "Immediate validation: one crop image in, one enhanced crop image out.")
    add_bullet(doc, "Final Phase 2 behavior: original image plus detected faces in, final blended high-resolution output image out.")
    add_bullet(doc, "Workflow JSON must not be edited manually during testing; runtime values are injected by backend code.")

    add_heading(doc, "3. Workflow Automation Strategy", 1)
    add_para(
        doc,
        "The exported ComfyUI API JSON will be stored as a workflow template. For every API request, the backend creates a deep copy of the template and injects request-specific values into known node inputs before submitting the payload to ComfyUI.",
    )
    node_rows = [
        ("958", "LoadImage", "inputs.image", "Uploaded cropped face image filename"),
        ("1056", "LoraLoaderModelOnly", "inputs.lora_name", "LoRA filename resolved from character_id"),
        ("1057", "LoraLoaderModelOnly", "inputs.lora_name", "Same character LoRA, unless workflow-specific mapping differs"),
        ("1071", "Text Multiline", "inputs.text", "Optional prompt or enhancement instruction"),
        ("866", "SaveImage", "inputs.filename_prefix", "Unique request_id or job_id output prefix"),
    ]
    add_simple_table(doc, ["Node ID", "Class", "Field", "Runtime Value"], node_rows, [900, 2100, 1850, 4510], LIGHT_GRAY, 9)

    add_heading(doc, "4. API Design", 1)
    add_heading(doc, "4.1 Minimum Crop Test Endpoint", 2)
    add_code(
        doc,
        """
POST /api/v1/face-enhance
Authorization: Bearer <api_key>
Content-Type: multipart/form-data

image: <cropped_face_image.jpg>
character_id: <character_or_lora_key>
prompt: <optional enhancement instruction>
""",
    )
    add_heading(doc, "4.2 Full Phase 2 Endpoint", 2)
    add_code(
        doc,
        """
POST /api/v1/face-swap
Authorization: Bearer <api_key>
Content-Type: multipart/form-data

image: <original_high_resolution_image.jpg>
faces: [
  {
    "face_id": "face_001",
    "bbox": {"xmin": 120, "ymin": 80, "width": 256, "height": 256},
    "character_id": "cousin_sean"
  }
]
""",
    )
    add_heading(doc, "4.3 Expected Response", 2)
    add_code(
        doc,
        """
{
  "job_id": "req_20260707_001",
  "status": "completed",
  "enhanced_crop_url": "https://.../outputs/req_20260707_001_crop.png",
  "final_image_url": "https://.../outputs/req_20260707_001_final.png",
  "metadata": {
    "workflow_template": "zooey.json",
    "comfyui_prompt_id": "...",
    "blend_method": "alpha_feathering"
  }
}
""",
    )

    add_heading(doc, "5. Four-Day Delivery Plan", 1)
    plan_rows = [
        (
            "Day 1",
            "Workflow extraction and one-crop ComfyUI test",
            "Confirm zooey.json is the API-format workflow template. Create node mapping. Upload one crop image to ComfyUI. Inject node 958, 1056, 1057, 1071, and 866. Submit /prompt and retrieve one enhanced crop output.",
            "Single crop test passes without manually editing JSON. Node mapping and test evidence are documented.",
        ),
        (
            "Day 2",
            "API wrapper, LoRA mapping, and account protection",
            "Build the Python ComfyUI client and protected POST endpoint. Add lora_config.json mapping character_id/name to exact LoRA filenames. Implement API key/account authentication and basic error responses.",
            "A caller can submit one crop image through API and receive an enhanced crop. Invalid API keys are rejected.",
        ),
        (
            "Day 3",
            "Multi-face loop, GCP hosting, and production-like test",
            "Deploy API and ComfyUI on GCP or a selected host platform. Add multi-character loop over detected faces. Confirm model, VAE, LoRA, and custom node availability. Test public API access with API key protection.",
            "Hosted endpoint works. Multiple detected faces can be processed sequentially. ComfyUI timeout and LoRA loading errors are handled.",
        ),
        (
            "Day 4",
            "Spatial merging, validation UI, and handoff",
            "Use Phase 1 bbox metadata to place enhanced faces back into the original image. Implement alpha feathering or Poisson blending. Extend the validation UI with upload, character assignment, execution logs, and comparison matrix.",
            "Final image shows no visible square crop boundary. Handoff includes API docs, deployment notes, test screenshots, and known limitations.",
        ),
    ]
    add_simple_table(doc, ["Day", "Focus", "Main Tasks", "Deliverables"], plan_rows, [800, 1900, 4350, 2310], LIGHT_BLUE, 8.3)

    add_heading(doc, "6. Day 1 Immediate Test Flow", 1)
    add_number(doc, "Prepare one cropped face image, for example test_crop_001.jpg.")
    add_number(doc, "Upload the crop to ComfyUI and capture the server-side image filename.")
    add_number(doc, "Load zooey.json and deep copy it into a request payload.")
    add_number(doc, "Inject the uploaded filename into node 958 inputs.image.")
    add_number(doc, "Resolve character_id to the correct LoRA filename and inject it into nodes 1056 and 1057.")
    add_number(doc, "Inject optional prompt text into node 1071.")
    add_number(doc, "Set node 866 filename_prefix to a unique job_id.")
    add_number(doc, "Submit the payload to ComfyUI /prompt and poll /history until the output is available.")
    add_number(doc, "Return the enhanced crop image and log request_id, prompt_id, LoRA, output path, and runtime.")

    add_heading(doc, "7. Spatial Merging and Blending Plan", 1)
    add_para(
        doc,
        "For the final Phase 2 output, the enhanced face cannot be pasted directly onto the original image. The backend must retrieve Phase 1 bounding-box metadata, resize the enhanced crop to the exact target region, and blend it into the source image using a soft mask.",
    )
    add_bullet(doc, "Required metadata: xmin, ymin, width, height, face_id, and character_id.")
    add_bullet(doc, "Initial method: PIL alpha feathering with Gaussian-blurred mask edges.")
    add_bullet(doc, "Higher-quality option: OpenCV seamlessClone / Poisson blending when lighting or skin tone mismatch is visible.")
    add_bullet(doc, "Acceptance standard: no visible square crop boundary and no harsh skin-tone discontinuity around the face region.")

    add_heading(doc, "8. Validation UI Expansion", 1)
    ui_rows = [
        ("Input Panel", "Image upload and manual character identity assignment for each detected face."),
        ("Execution Logs", "Real-time states such as Queueing, Processing, ComfyUI completed, Blending, and Failed."),
        ("Comparison Matrix", "Side-by-side view of original image, isolated ComfyUI face output, and final merged image."),
        ("Debug Metadata", "Show face_id, bbox, selected LoRA, ComfyUI prompt_id, and output file path for validation."),
    ]
    add_simple_table(doc, ["UI Area", "Requirement"], ui_rows, [2200, 7160], LIGHT_GRAY, 9.2)

    add_heading(doc, "9. Risks and Mitigation", 1)
    risk_rows = [
        ("MediaPipe/face detection misses faces", "Production must not skip faces", "Keep InsightFace or a hybrid detector as fallback and process all Phase 1 detected regions."),
        ("Workflow node IDs change", "Runtime injection fails", "Store node mapping in config and validate required nodes at service startup."),
        ("LoRA file missing or fails to load", "ComfyUI job fails", "Validate lora_config.json against server files and return structured errors."),
        ("ComfyUI timeout or worker crash", "API request hangs or fails", "Set timeout, job status polling, retry policy, and clear failure response."),
        ("Poor crop boundary blending", "Final image looks visibly edited", "Use feathered masks first, then Poisson blending for difficult lighting cases."),
        ("Public endpoint abuse", "Security and cost risk", "Expose only the API service, require API key/account, and keep ComfyUI UI private."),
    ]
    add_simple_table(doc, ["Risk", "Impact", "Mitigation"], risk_rows, [2400, 2600, 4360], LIGHT_GRAY, 8.8)

    add_heading(doc, "10. Acceptance Criteria", 1)
    add_bullet(doc, "A single protected POST endpoint can automate the Phase 1 + Phase 2 chain.")
    add_bullet(doc, "The crop-only test returns an enhanced cropped face image through the API.")
    add_bullet(doc, "The backend dynamically injects crop image and LoRA filename into the ComfyUI JSON payload.")
    add_bullet(doc, "Character-to-LoRA mapping is stored in lora_config.json or equivalent configuration.")
    add_bullet(doc, "Multiple detected faces can be processed one by one, with the correct LoRA for each character.")
    add_bullet(doc, "Enhanced faces are merged back into original coordinates without visible square boundaries.")
    add_bullet(doc, "Validation UI shows original image, isolated ComfyUI output, final merged image, and execution logs.")
    add_bullet(doc, "Timeouts, missing LoRA files, and ComfyUI failures do not crash the server.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
