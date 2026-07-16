from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("output/docs/AIFX_Phase2_Face_Enhancement_API_4_Day_Plan_CN.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 32, 46)
MUTED = RGBColor(90, 96, 106)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.08


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_run_font(run, size=None, color=None, bold=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def add_para(doc, text="", style=None, size=11, color=INK, bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)
        r = p.add_run(line)
        set_run_font(r, size=9.5, color=INK, name="Courier New")


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_margins(table, 120, 160, 120, 160)
    set_table_borders(table, "C8D5E2")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    add_para(doc, "", after=6)


def add_simple_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_cell_margins(table)
    set_table_borders(table)
    set_repeat_header(table.rows[0])
    hdr = table.rows[0].cells
    for idx, label in enumerate(headers):
        set_cell_shading(hdr[idx], header_fill)
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=font_size, color=INK, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(value)
            set_run_font(r, size=font_size, color=INK)
    return table


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = "AIFX Studio Phase 2 | Face Enhancement API Plan"
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Confidential working plan"
    set_run_font(footer.runs[0], size=9, color=MUTED)

    add_para(doc, "AIFX Studio Phase 2", size=13, color=MUTED, bold=True, after=4)
    add_para(doc, "Face Enhancement API 四天交付计划书", size=24, color=INK, bold=True, after=4)
    add_para(
        doc,
        "目标：用 1 张 crop face image 调用 ComfyUI workflow，返回 enhanced crop 测试结果，并为后续合成回原图预留接口。",
        size=12,
        color=MUTED,
        after=14,
    )

    meta_rows = [
        ("Prepared for", "Phase 2 implementation / API-based testing"),
        ("Prepared by", "Zooey Chen"),
        ("Date", "2026-07-07"),
        ("Primary workflow", "ComfyUI API prompt template: zooey.json"),
        ("Core test case", "Input one cropped face image, output one enhanced cropped face image"),
    ]
    for label, value in meta_rows:
        p = add_para(doc, "", after=2)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=10.5, color=INK, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=10.5, color=INK)

    add_callout(
        doc,
        "Executive Summary / 核心结论",
        "本阶段不应手动修改 ComfyUI 导出的 JSON。zooey.json 将作为固定 workflow template 保存，API backend 每次收到请求时复制模板，并自动注入输入图片、LoRA、prompt 和输出前缀，然后提交到 ComfyUI /prompt。四天内优先完成 crop image 输入到 enhanced crop 输出的 API 闭环。",
    )

    add_heading(doc, "1. Project Objective / 项目目标", 1)
    add_para(
        doc,
        "Phase 2 的目标是把当前 ComfyUI face enhancement workflow 封装成可被外部系统调用的 API 服务。测试时，调用方会通过 API token 或 API account 提交图片，服务端自动调用 ComfyUI 并返回结果。",
    )
    add_bullet(doc, "最小可交付闭环：输入 1 张 cropped face image，输出 1 张 enhanced cropped face image。")
    add_bullet(doc, "后续扩展闭环：输入 original image、crop bbox 和 crop image，输出合成回原图后的 final image。")
    add_bullet(doc, "领导要求：不能手动修改 workflow JSON；必须由代码在运行时自动注入参数。")

    add_heading(doc, "2. Confirmed Scope / 当前范围", 1)
    scope_rows = [
        ("P0", "Crop image enhancement", "API 接收 1 张 crop face image，调用 ComfyUI，返回 enhanced crop。", "Day 1-2"),
        ("P1", "API account access", "通过 API key / account 限制访问，避免公开无保护访问。", "Day 2"),
        ("P2", "GCP hosting", "部署 API 和 ComfyUI 到 GCP，解决 host platform 交付问题。", "Day 3"),
        ("P3", "Merge back to original", "使用 bbox + feather / Gaussian blur mask 把 enhanced crop 合成回原图。", "Day 4 or follow-up"),
    ]
    add_simple_table(doc, ["Priority", "Module", "Description", "Timing"], scope_rows, [950, 2100, 4800, 1510], LIGHT_BLUE)

    add_heading(doc, "3. Workflow Template Strategy / JSON 自动注入策略", 1)
    add_para(
        doc,
        "zooey.json 已经是 ComfyUI API prompt 格式，可以作为后端固定模板。API 不直接让用户修改 JSON，而是在每次请求时 deep copy 该模板，并只替换以下节点字段。",
    )
    node_rows = [
        ("958", "LoadImage", "inputs.image", "请求上传后的 crop image filename"),
        ("1056", "LoraLoaderModelOnly", "inputs.lora_name", "根据 character_id 自动选择 LoRA"),
        ("1057", "LoraLoaderModelOnly", "inputs.lora_name", "保持与 1056 一致，或按 workflow 分别配置"),
        ("1071", "Text Multiline", "inputs.text", "额外 prompt，例如 clothing / mask / style instruction"),
        ("866", "SaveImage", "inputs.filename_prefix", "使用 request_id / job_id 生成唯一输出路径"),
    ]
    add_simple_table(doc, ["Node ID", "Class", "Field", "Runtime Value"], node_rows, [900, 2050, 1850, 4560], LIGHT_GRAY, 9.3)

    add_callout(
        doc,
        "Implementation Rule / 实现规则",
        "workflow JSON 本身作为 template 只读保存。所有图片名、LoRA、prompt、输出路径都由 API backend 自动写入临时 payload。若 ComfyUI workflow 未来重导出，只需更新 node mapping 配置，不应在测试时手动改 JSON。",
    )

    add_heading(doc, "4. API Input and Output / API 输入输出", 1)
    add_heading(doc, "4.1 Minimum Test Endpoint", 2)
    add_code(
        doc,
        """
POST /api/v1/face-enhance
Authorization: Bearer <api_key>
Content-Type: multipart/form-data

image: <crop_face_image.jpg>
character_id: <character_or_lora_key>
prompt: <optional text instruction>
""",
    )
    add_heading(doc, "4.2 Expected Response", 2)
    add_code(
        doc,
        """
{
  "job_id": "req_20260707_001",
  "status": "completed",
  "enhanced_crop_url": "https://.../outputs/req_20260707_001.png",
  "metadata": {
    "workflow_template": "zooey.json",
    "character_id": "cousin_sean",
    "comfyui_prompt_id": "..."
  }
}
""",
    )

    add_heading(doc, "5. Four-Day Delivery Plan / 四天完成计划", 1)
    plan_rows = [
        (
            "Day 1",
            "Workflow and single-image test",
            "固定 zooey.json 作为 template；确认 node mapping；上传 1 张 crop image 到 ComfyUI；自动注入 node 958 / 1056 / 1057 / 1071 / 866；POST /prompt 并取回结果。",
            "本地单张 crop image 测试成功；产出 node mapping 表和 API I/O 草案。",
        ),
        (
            "Day 2",
            "API service and account access",
            "实现 POST /api/v1/face-enhance；支持 multipart image upload；增加 API key / account 校验；封装 ComfyUI upload、prompt submit、history polling 和 output retrieval。",
            "可用 API 调用，不需要手动改 JSON；API key 错误会拒绝访问。",
        ),
        (
            "Day 3",
            "GCP deployment and host platform",
            "在 GCP 部署 ComfyUI 与 API 服务；配置模型、VAE、LoRA、custom nodes；开放受保护 API endpoint；从本地调用云端接口测试。",
            "GCP public API endpoint 可访问；通过 API key 上传 crop image 并返回 enhanced crop。",
        ),
        (
            "Day 4",
            "Stability, merge preparation, and handoff",
            "补充错误处理、日志、timeout、job status；测试多张 crop；准备 feather / Gaussian blur merge-back 方案；整理 README、测试截图、接口说明和已知限制。",
            "最终 demo、接口文档、测试结果和 Phase 2 handoff summary。",
        ),
    ]
    add_simple_table(doc, ["Day", "Focus", "Main Tasks", "Deliverables"], plan_rows, [850, 1850, 4300, 2360], LIGHT_BLUE, 8.8)

    add_heading(doc, "6. Day 1 Immediate Execution Flow / 第一天立即执行流程", 1)
    add_number(doc, "准备 1 张 crop face image，文件名保持简单，例如 test_crop_001.jpg。")
    add_number(doc, "通过 ComfyUI upload API 上传图片，取得 ComfyUI 可识别的 filename。")
    add_number(doc, "读取 zooey.json template，并 deep copy 成本次 request payload。")
    add_number(doc, "把 node 958 的 inputs.image 改为上传后的 crop image filename。")
    add_number(doc, "把 node 1056 和 1057 的 inputs.lora_name 改为目标角色 LoRA 文件。")
    add_number(doc, "把 node 1071 的 inputs.text 改为可选 prompt；若无 prompt，使用空字符串或默认增强描述。")
    add_number(doc, "把 node 866 的 filename_prefix 改为 request_id，避免输出文件互相覆盖。")
    add_number(doc, "POST payload 到 ComfyUI /prompt，并轮询 /history 获取输出。")
    add_number(doc, "返回 enhanced crop image，记录输入、输出、prompt_id 和耗时。")

    add_heading(doc, "7. GCP Hosting Direction / GCP 部署方向", 1)
    add_para(
        doc,
        "如果需要云端部署，建议把 API service 和 ComfyUI 放在同一台 GCP GPU VM 或同一私有网络内，减少图片传输和权限复杂度。公网只暴露 API service，不直接暴露 ComfyUI 管理端口。",
    )
    add_bullet(doc, "Public access: 只开放 API endpoint，并要求 Authorization header。")
    add_bullet(doc, "Private ComfyUI: ComfyUI 只允许 API service 调用，避免任何人直接访问 workflow UI。")
    add_bullet(doc, "Storage: 输入 crop、输出 enhanced crop 可先存本地磁盘，后续再接 Cloud Storage。")
    add_bullet(doc, "Security: API key 存在环境变量或数据库，不写死在代码和 Git repo。")

    add_heading(doc, "8. Risks and Mitigation / 风险与处理", 1)
    risk_rows = [
        ("ComfyUI workflow node 变化", "重新导出 workflow 后 node id 可能变化", "维护 workflow_mapping 配置；启动时校验关键节点是否存在"),
        ("LoRA 或模型缺失", "云端环境可能没有本地模型文件", "Day 3 前列出完整模型清单，并部署后做健康检查"),
        ("ComfyUI job timeout", "图片大或 GPU 忙时等待过久", "设置 timeout、job status、错误响应和日志"),
        ("Crop 边缘融合问题", "enhanced crop 回贴原图可能出现边缘痕迹", "Day 4 准备 feather / Gaussian blur mask blend 方案"),
        ("公开访问安全", "服务被无权限调用或滥用", "API account / API key；限制 ComfyUI 端口；记录 request log"),
    ]
    add_simple_table(doc, ["Risk", "Impact", "Mitigation"], risk_rows, [2500, 3100, 3760], LIGHT_GRAY, 9)

    add_heading(doc, "9. Acceptance Criteria / 验收标准", 1)
    add_bullet(doc, "调用方可以通过 API 上传 1 张 crop face image。")
    add_bullet(doc, "服务端自动调用 zooey.json 对应的 ComfyUI workflow。")
    add_bullet(doc, "测试过程不需要手动修改 workflow JSON。")
    add_bullet(doc, "API 返回 enhanced crop image 或可访问的 output URL。")
    add_bullet(doc, "API 受 API account / API key 保护，未授权请求会被拒绝。")
    add_bullet(doc, "文档包含 API input/output、部署方式、测试结果和已知限制。")

    add_heading(doc, "10. Suggested Final Deliverables / 最终交付物", 1)
    add_bullet(doc, "Source code repository with API service and ComfyUI client wrapper.")
    add_bullet(doc, "zooey.json workflow template and workflow_mapping configuration.")
    add_bullet(doc, "GCP deployment notes and environment variable example.")
    add_bullet(doc, "API usage document with request / response examples.")
    add_bullet(doc, "Single crop image test result and screenshots.")
    add_bullet(doc, "Optional merge-back prototype for enhanced crop to original image.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
